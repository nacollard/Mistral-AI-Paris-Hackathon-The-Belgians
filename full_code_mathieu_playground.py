from mistralai.client import MistralClient
from mistralai.models.chat_completion import ChatMessage
from dotenv import load_dotenv
from OpenRAG.src.openrag.chunk_vectorization.chunk_vectorization import get_vectorizer
from OpenRAG.src.openrag.vectordb.milvus_adapter import init_milvus_connection
from pymilvus import Collection
from docx import Document
from requests.exceptions import HTTPError, Timeout, ConnectionError
import os
import json
import re
import requests
import xml.etree.ElementTree as ET

load_dotenv()  # take environment variables from .env.


def get_xml_element_text(xml_response, tag, default=None):
    """
    Retrieve the content between the specified XML tags.

    Args:
        xml_response (str): The XML response.
        tag (str): The XML tag to retrieve the content from.
        default (optional): The value to return if the tag is not found.

    Returns:
        str or default: The content between the specified XML tags or the default value.
    """
    try:
        root = ET.fromstring(xml_response)
        element = root.find(f"./{tag}")
        if element is not None:
            return element.text
        else:
            return default
    except ET.ParseError:
        return default

def send_request_to_mistral_ai_model(model, messages):
    """
    Send a request to the Mistral AI model and return the response.

    Args:
        model (str): The Mistral AI model to use.
        messages (List[ChatMessage]): The messages to send in the request.

    Returns:
        str: The response from the Mistral AI model.
    """
    api_key = os.environ["MISTRAL_API_KEY"]
    if not api_key:
        raise ValueError("MISTRAL_API_KEY not found in environment variables.")
    client = MistralClient(api_key=api_key)

    try:
        chat_response = client.chat(
            model=model,
            messages=messages,
        )

        response = chat_response.choices[0].message.content
        return response
    except HTTPError as e:
        print(f"Error in Mistral AI request: HTTPError ({e})")
        return None
    except Timeout as e:
        print(f"Error in Mistral AI request: Timeout ({e})")
        return None
    except ConnectionError as e:
        print(f"Error in Mistral AI request: ConnectionError ({e})")
        return None
    except Exception as e:
        print(f"Error in Mistral AI request: Unexpected error ({e})")
        return None


def load_company_knowledge():
    """
    Load company knowledge from company documents.

    Returns:
        str: The combined text from all company documents.
    """
    business_model = "Data/Internal/Business Model de StIT.docx"
    long_term_strategy = (
        "Data/Internal/Plan de développement stratégique sur 8 ans pour StIT.docx"
    )
    products_and_services = "Data/Internal/Produits et services de StIT.docx"
    company_docs = [business_model, long_term_strategy, products_and_services]
    company_knowledge = ""

    for document_path in company_docs:
        docx_document = Document(document_path)
        paragraphs_text = " ".join(
            [paragraph.text for paragraph in docx_document.paragraphs]
        )
        company_knowledge += paragraphs_text

    return company_knowledge


def find_chunks(id, path="Data/Internal/HR/"):
    """
    Find the chunk based on the given id.

    Args:
        id (int): The id of the chunk to find.
        path (str, optional): The path to the chunk files. Defaults to "Data/Internal/HR/".

    Returns:
        dict: A dictionary containing the chunk details, or None if the chunk is not found.
    """
    global_indexing = json.load(open("global_indexing.json", "r"))
    for key, value in global_indexing.items():
        start_idx = value["start"]
        end_idx = value["end"]
        if start_idx <= id <= end_idx:
            index_in_file = id - start_idx
            data_dict_file = json.load(open(path + key + "_chunks.json", "r"))
            data_dict_file["chunk_" + str(index_in_file)]["document"] = key + ".docx"
            data_dict_file["chunk_" + str(index_in_file)]["fullpath"] = (
                path + key + ".docx"
            )
            return data_dict_file["chunk_" + str(index_in_file)]
    return None


def create_prompt_analyst_agent(context, company_knowledge, type):
    """
    Create the prompt for the Mistral AI model.

    Args:
        article (str): The news article to analyze.
        company_knowledge (str): The company knowledge to include in the prompt.

    Returns:
        str: The prompt for the Mistral AI model.
    """
    prompt = f"""
                You are an experienced business analyst tasked with determining the priority level of {type}s based on their relevance to your company, StIT.

                Here is some crucial information about the company to consider during your analysis:
                <company_knowledge>{company_knowledge}</company_knowledge>

                Please thoroughly read and analyze the following {type}:

                <{type}>{context}</{type}>

                After completing your analysis, provide your final assessment in the <output> section, using the following format:

                <output>
                <priority_level>High OR Medium OR Low</priority_level>
                <justification>A detailed explanation of your priority rating, including how the {type}'s main points and key details relate to the company's goals, operations, or industry, and the potential implications and impact of the {type} on the company</justification>
                <main_topic>A one-sentence summary highlighting the {type}'s main topic<main_topic>
                </output>

                Remember, your goal is to help company management quickly identify and prioritize important {type}s, so be sure to consider the key implications and potential impact of the {type} on the company in your priority rating and justification.
            """
    return prompt


def create_prompt_strategy_agent(context, stretegic_info):
    prompt = f"""
                You are an experienced strategic consultant.
                I will provide you with some context about a situation and the employee you have to advice. I'll also provide you strategic information about the company.
                Your task is to build a quick action plan for the employee the message is addressed to, on how they can best tackle the matter to help the company based on their experience and position in the company.

                Here is the context with the employee you have to advice and information about the situation:
                <context>{context}</context>

                Here is the strategic information about the company:
                <strategic_info>{stretegic_info}</strategic_info>

                First, take a moment to carefully read and understand the context, strategic information, and who is the employee you advice. Think through the key considerations and how the employee can best respond given the company's strategic priorities and his or her position. Write your thoughts in a <scratchpad> section.

                Then, provide a short action plan with a few concrete steps the employee can follow to effectively address this situation in a way that aligns with and supports the company's strategy. Write the action plan inside <action_plan> tags.

                The action plan should be concise and to-the-point, focusing on the most critical steps the employee should take. Aim for 3-5 key action items.

                Remember, your goal is to guide the employee on how to tackle this matter in a way that will best help the company achieve its strategic objectives, based on your understanding of the context, company strategy, and the specific situation described in the message.
                """
    return prompt


def strategy_agent(context, employee):
    strategic_info = company_information
    cv_path = Document(f"Data/Internal/HR/CV {employee}.docx")
    employee_CV = " ".join([paragraph.text for paragraph in cv_path.paragraphs])
    strategic_info += employee_CV
    prompt = create_prompt_strategy_agent(context, strategic_info)
    messages = [
        ChatMessage(role="system", content=prompt),
        ChatMessage(role="user", content="The matter : " + context),
    ]

    model = "mistral-large-latest"

    response = send_request_to_mistral_ai_model(model, messages)
    print("Strategic advisor response: ", response)
    return response


def analyst_agent(context, type):
    """
    Analyze the given news article and dispatch it to the appropriate agent.

    Args:
        article (str): The news article to analyze.

    Returns:
        None
    """
    context_content = context

    company_knowledge = company_information

    prompt = create_prompt_analyst_agent(context_content, company_knowledge, type)

    model = "mistral-large-latest"

    messages = [
        ChatMessage(role="system", content=prompt),
        ChatMessage(
            role="user", content=f"Content of the {type}  : " + context_content
        ),
    ]

    xml_response = send_request_to_mistral_ai_model(model, messages)

    priority_level = get_xml_element_text(xml_response, "priority_level")
    justification = get_xml_element_text(xml_response, "justification")
    main_topic = get_xml_element_text(xml_response, "main_topic")

    print("Priority Level: ", priority_level)
    print("Justification: ", justification)
    print("Main Topic: ", main_topic)

    employees_to_inform = dispatch_agent(main_topic, justification)
    strategy = "No strategy"
    if priority_level == "High":
         print("High priority level detected.")
         employee = employees_to_inform[0]
         print(f"/!\ Informing {employee} about the high priority {type}. /!\ ")
         context_to_pass = f"<employee_name>employee</employee_name>, <priority_level>High</priority_level>, <main_topic>{main_topic}</main_topic>, <justification>{justification}</justification>, <context>{context}</context>, <type>{type}</type>"
         strategy = strategy_agent(context_to_pass, employee)
    return employees_to_inform, priority_level, main_topic, context, justification, type, strategy


def create_prompt_dispatch_agent(main_topic, justification, CVs):
    """
    Create the prompt for the Mistral AI model.

    Args:
        maint_topic (str): The main topic of the news article to analyze.
        justification (str): The reason why the news article might be relevant to the company.
        CVs (str): The CVs of the employees.

    Returns:
        str: The prompt for the Mistral AI model.
    """

    prompt = f""" You are a senior executive at StIT. Your task is to identify and rank employees based on their relevance to a specific matter, considering their expertise and roles.
                Read and analyze the matter:
                <matter>{main_topic} {justification}</matter>

                Next, review the CVs of the employees:
                <CVs>{CVs}</CVs>

                In the <output> section, list the names of the relevant employees in order of relevance. Ensure that the main points and key details of the matter align with the provided CVs and job titles at StIT.

                Use the following format for your final assessment:
                <output>
                <total>NUMBER</total>
                <employee1>NAME_EMPLOYEE1</employee1>
                <employee2>NAME_EMPLOYEE2</employee2>
                <employee3>NAME_EMPLOYEE3</employee3>
                <employee4>NAME_EMPLOYEE4</employee4>
                <employee5>NAME_EMPLOYEE5</employee5>
                </output>

                In the <total> tag, specify the number of employees to be informed. In the <employeeN> tags, provide the names of the selected employees. If no employee is relevant, write "None" in the <total> tag.

                Your goal is to help management quickly identify and prioritize employees to inform. Consider the implications and potential impact of the matter on the company. Remember, employees receive push notifications, so only include those who are truly relevant.
                Do not provide any additional information or context beyond the total employees to inform and their names. Do not justify your choices, only give back the output section previously defined.
                """
    return prompt


def dispatch_agent(main_topic, justification):
    """
    Determine the employee(s) who should be informed about the given topic and justification.

    Args:
        main_topic (str): The main topic of the news.
        justification (str): The justification for the priority level of the news.

    Returns:
        None
    """
    results = internal_retriever_agent(main_topic)

    CVs = []
    for result in results:
        content_cv = ""
        docx_document = Document(result["fullpath"])
        paragraphs_text = " ".join(
            [paragraph.text for paragraph in docx_document.paragraphs]
        )
        content_cv += paragraphs_text
        # The input string
        filename = result["fullpath"]

        # The regular expression pattern to match the name
        pattern = r"/CV\s*(.+?)\.docx"

        # Search for the pattern in the input string
        match = re.search(pattern, filename)
        name = match.group(1)
        # The name is in between 'CV' and '.docx'
        CVs.append("Name: " + name + " " + content_cv)

    prompt = create_prompt_dispatch_agent(main_topic, justification, CVs)
    messages = [
        ChatMessage(role="system", content=prompt),
        ChatMessage(role="user", content="The matter : " + main_topic + justification),
    ]

    model = "mistral-large-latest"

    employees_to_inform = send_request_to_mistral_ai_model(model, messages)
    pattern = r"<output>.*</output>"
    print("Employees to inform: ", employees_to_inform)
    try:
        # Search for the pattern in the string and extract the match
        match = re.search(pattern, employees_to_inform, re.DOTALL)
        xml_content = match.group(0) if match else None
        # Parse the XML string into an ElementTree object
        root = ET.fromstring(xml_content.strip())
    except re.error as e:
        # Handle the exceptions raised by the re module
        print(employees_to_inform)
        print(f"An error occurred while processing the regular expression: {e}")
    except ET.ParseError as e:
        # Handle the exceptions raised by the ET module
        print(xml_content)
        print(f"An error occurred while parsing the XML string: {e}")
    except Exception as e:
        # Handle all the other exceptions
        print(f"An unexpected error occurred: {e}")

    # Define an empty list to store the employee names
    number_of_employees_to_warn = 0
    employee_names = []

    # Iterate over the child elements of the root element
    for child in root:
        if child.tag == "total":
            number_of_employees_to_warn = int(child.text)
        # If the child element is an employee element, extract the name and append it to the list
        if child.tag.startswith("employee"):
            employee_names.append(child.text)

    return employee_names


def internal_retriever_agent(text, filter="HR"):
    """
    Retrieve internal information related to the given text.

    Args:
        text (str): The text to search for.
        filter (str, optional): The filter to apply to the search results. Defaults to 'HR'.

    Returns:
        list: A list of dictionaries containing the search results.
    """
    vectorizer = get_vectorizer("mistral")
    query_vector = vectorizer.vectorize(text)

    init_milvus_connection()

    collection_name = "mistral_collection"
    collection = Collection(name=collection_name)

    n_neighbors = 20
    results = collection.search(
        [query_vector],
        "vector",
        param={"metric_type": "L2", "params": {}},
        limit=n_neighbors,
        expr="source == '" + filter + "'",
    )

    final_indices = []
    for result in results[0]:
        if result.id not in final_indices and len(final_indices) + 1 <= n_neighbors:
            final_indices.append([result.id, result.distance])
            prev_index = result.id - 1
            if (
                prev_index >= 0
                and prev_index not in final_indices
                and len(final_indices) + 1 <= n_neighbors
                and filter != "HR"
            ):
                final_indices.append([prev_index, result.distance])
            next_index = result.id + 1
            if (
                next_index not in final_indices
                and len(final_indices) + 1 <= n_neighbors
                and filter != "HR"
            ):
                final_indices.append([next_index, result.distance])

    results = final_indices

    answer_chunks = []
    unique_chunks = []
    answer_chunks2 = []
    for hit in results:
        answer_chunk = find_chunks(hit[0])
        if answer_chunk["document"] in unique_chunks and filter == "HR":
            continue
        unique_chunks.append(answer_chunk["document"])
        answer_chunks.append(answer_chunk["text"])
        answer_chunks2.append(answer_chunk)

    return answer_chunks2


news_article1 = 'Data/External/News Articles/ New tax law in France aims to encourage and support the growth of startups and small businesses copy.docx' 
news_article2 = 'Data/External/News Articles/Local bakery in Paris wins award for best croissant in the city copy.docx' 
news_article3 = 'Data/External/News Articles/New survey finds that the majority of French people prefer to shop online rather than in-store copy.docx' 
news_article4 = 'Data/External/News Articles/Massive cyberattack exposes the vulnerabilities of businesses and organizations copy.docx' 

new_law1 = 'Data/External/Laws/New animal welfare law in France.docx' 
new_law2 = 'Data/External/Laws/New data privacy law in France.docx' 
new_law3 = 'Data/External/Laws/New labor and employment law in France.docx' 
new_law4 = 'Data/External/Laws/New cybersecurity and technology law in France.docx' 
            
new_post1 = 'Data/External/Social Media/cybersecurity.docx' 
new_post2 = 'Data/External/Social Media/space_travel.docx' 
new_post3 = 'Data/External/Social Media/startup_culture.docx' 
new_post4 = 'Data/External/Social Media/coffee_and_snacks_lover.docx' 
            
company_information = load_company_knowledge()

if __name__ == "__main__": 
    print("######################")
    print("New Incoming Message: ")
    print("######################")
    analyst_agent(news_article2, "news article") 
    #analyst_agent(news_article4, "news article") 
    print("######################")
    print("New Incoming Message: ")
    print("######################")
    analyst_agent(new_law3, "law") 
    #analyst_agent(new_law4, "law") 
    print("######################")
    print("New Incoming Message: ")
    print("######################")
    analyst_agent(new_post1, "social media post") 
    #analyst_agent(new_post4, "social media post")